import pandas as pd
import json
from typing import Union, BinaryIO, Dict, Any
from pathlib import Path
import docx
import requests
from PIL import Image
import io
import openpyxl
import urllib.parse

class FilePreprocessor:
    """Handles preprocessing of various file types for data processing pipeline."""
    
    @staticmethod
    def process_excel(file: Union[BinaryIO, str]) -> pd.DataFrame:
        """
        Process Excel files (.xlsx) and convert to pandas DataFrame
        
        Args:
            file: File object or path to Excel file
            
        Returns:
            pd.DataFrame: Processed data as DataFrame
        """
        try:
            return pd.read_excel(file)
        except Exception as e:
            raise ValueError(f"Error processing Excel file: {str(e)}")

    @staticmethod
    def process_csv(file: Union[BinaryIO, str]) -> pd.DataFrame:
        """
        Process CSV files and convert to pandas DataFrame
        
        Args:
            file: File object or path to CSV file
            
        Returns:
            pd.DataFrame: Processed data as DataFrame
        """
        try:
            return pd.read_csv(file)
        except Exception as e:
            raise ValueError(f"Error processing CSV file: {str(e)}")

    @staticmethod
    def process_json(file: Union[BinaryIO, str]) -> str:
        """
        Process JSON files and convert to string
        
        Args:
            file: File object or path to JSON file
            
        Returns:
            str: JSON content as string
        """
        try:
            if isinstance(file, str):
                with open(file, 'r') as f:
                    data = json.load(f)
            else:
                data = json.load(file)
            return json.dumps(data)
        except Exception as e:
            raise ValueError(f"Error processing JSON file: {str(e)}")

    @staticmethod
    def process_text(file: Union[BinaryIO, str]) -> str:
        """
        Process text files (.txt) and convert to string
        
        Args:
            file: File object or path to text file
            
        Returns:
            str: File content as string
        """
        try:
            if isinstance(file, str):
                with open(file, 'r') as f:
                    return f.read()
            return file.read().decode('utf-8')
        except Exception as e:
            raise ValueError(f"Error processing text file: {str(e)}")

    @staticmethod
    def process_docx(file: Union[BinaryIO, str]) -> str:
        """
        Process Word documents (.docx) and convert to string
        
        Args:
            file: File object or path to Word document
            
        Returns:
            str: Document content as string
        """
        try:
            if isinstance(file, str):
                doc = docx.Document(file)
            else:
                doc = docx.Document(io.BytesIO(file.read()))
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise ValueError(f"Error processing Word document: {str(e)}")

    @staticmethod
    def process_image(file: Union[BinaryIO, str], output_path: str = None) -> str:
        """
        Process image files (.png) and convert to JPEG
        
        Args:
            file: File object or path to image file
            output_path: Optional path to save converted image
            
        Returns:
            str: Path to converted JPEG file
        """
        try:
            if isinstance(file, str):
                img = Image.open(file)
            else:
                img = Image.open(io.BytesIO(file.read()))
            
            # Convert to RGB if necessary (PNG might have RGBA)
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            
            if output_path:
                jpeg_path = output_path
            else:
                # Generate output path if not provided
                original_path = file if isinstance(file, str) else 'image'
                jpeg_path = f"{Path(original_path).stem}.jpeg"
            
            img.save(jpeg_path, 'JPEG')
            return jpeg_path
        except Exception as e:
            raise ValueError(f"Error processing image: {str(e)}")

    @staticmethod
    def process_web_url(url: str) -> pd.DataFrame:
        """
        Process web URLs (Google Sheets or Excel) and convert to DataFrame
        
        Args:
            url: URL to the web sheet
            
        Returns:
            pd.DataFrame: Processed data as DataFrame
        """
        try:
            # Handle Google Sheets URLs
            if 'docs.google.com/spreadsheets' in url:
                # Convert to export URL
                file_id = url.split('/d/')[1].split('/')[0]
                export_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=csv"
                return pd.read_csv(export_url)
            
            # Handle direct Excel/CSV URLs
            response = requests.get(url)
            if response.status_code != 200:
                raise ValueError(f"Failed to fetch URL: {url}")
            
            content_type = response.headers.get('content-type', '')
            if 'csv' in content_type:
                return pd.read_csv(io.StringIO(response.text))
            elif 'excel' in content_type or 'spreadsheet' in content_type:
                return pd.read_excel(io.BytesIO(response.content))
            else:
                raise ValueError(f"Unsupported content type: {content_type}")
        except Exception as e:
            raise ValueError(f"Error processing web URL: {str(e)}")

    @classmethod
    def preprocess_file(cls, file: Union[BinaryIO, str], file_type: str, **kwargs) -> Union[pd.DataFrame, str]:
        """
        Main method to preprocess files based on their type
        
        Args:
            file: File object or path to file
            file_type: Type of file (e.g., 'xlsx', 'csv', 'json', etc.)
            **kwargs: Additional arguments for specific processors
            
        Returns:
            Union[pd.DataFrame, str]: Processed data
        """
        processors = {
            'xlsx': cls.process_excel,
            'csv': cls.process_csv,
            'json': cls.process_json,
            'txt': cls.process_text,
            'docx': cls.process_docx,
            'png': cls.process_image,
            'web_url': cls.process_web_url
        }
        
        processor = processors.get(file_type.lower())
        if not processor:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return processor(file, **kwargs)